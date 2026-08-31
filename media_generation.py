from __future__ import annotations

import base64
import io
import re

import httpx

TIMEOUT = 60.0
VALID_VOICES = {"alloy", "echo", "fable", "onyx", "nova", "shimmer"}
VALID_FORMATS = {"md", "txt", "pdf", "docx"}
CONTENT_TYPES = {
    "md": "text/markdown",
    "txt": "text/plain",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


async def generate_audio(text: str, base_url: str, api_key: str, voice: str = "alloy") -> bytes:
    if voice not in VALID_VOICES:
        voice = "alloy"
    url = base_url.rstrip("/") + "/audio/speech"
    headers = {"Authorization": f"Bearer {api_key}"}
    payload = {"model": "tts-1", "input": text, "voice": voice}
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            r = await client.post(url, headers=headers, json=payload)
        r.raise_for_status()
    except httpx.HTTPStatusError as e:
        raise RuntimeError(f"audio generation failed ({e.response.status_code}): {e.response.text[:300]}") from e
    except httpx.HTTPError as e:
        raise RuntimeError(f"audio generation request failed: {e}") from e
    return r.content


async def generate_image(prompt: str, base_url: str, api_key: str, size: str = "1024x1024") -> bytes:
    url = base_url.rstrip("/") + "/images/generations"
    headers = {"Authorization": f"Bearer {api_key}"}
    payload = {"model": "dall-e-3", "prompt": prompt, "n": 1, "size": size, "response_format": "b64_json"}
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            r = await client.post(url, headers=headers, json=payload)
        r.raise_for_status()
    except httpx.HTTPStatusError as e:
        raise RuntimeError(f"image generation failed ({e.response.status_code}): {e.response.text[:300]}") from e
    except httpx.HTTPError as e:
        raise RuntimeError(f"image generation request failed: {e}") from e

    data = r.json()
    try:
        b64 = data["data"][0]["b64_json"]
    except (KeyError, IndexError) as e:
        raise RuntimeError(f"unexpected response shape from image API: {data}") from e
    return base64.b64decode(b64)


def _plain_bytes(content: str) -> bytes:
    return content.encode("utf-8")


def _to_pdf(content: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    for line in content.splitlines() or [""]:
        heading = re.match(r"^(#{1,3})\s+(.*)", line)
        if heading:
            pdf.set_font("Helvetica", "B", 16 - 2 * (len(heading.group(1)) - 1))
            pdf.multi_cell(0, 8, heading.group(2))
            pdf.set_font("Helvetica", size=11)
        elif not line.strip():
            pdf.ln(6)
        else:
            text = re.sub(r"^[-*]\s+", "-  ", line)
            safe_text = text.encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 6, safe_text)
    return bytes(pdf.output())


def _to_docx(content: str) -> bytes:
    from docx import Document

    doc = Document()
    for line in content.splitlines() or [""]:
        heading = re.match(r"^(#{1,3})\s+(.*)", line)
        if heading:
            doc.add_heading(heading.group(2), level=len(heading.group(1)))
        elif re.match(r"^[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_document(content: str, fmt: str) -> bytes:
    if fmt not in VALID_FORMATS:
        raise ValueError(f"unsupported format: {fmt}")
    if fmt in ("md", "txt"):
        return _plain_bytes(content)
    if fmt == "pdf":
        return _to_pdf(content)
    return _to_docx(content)
