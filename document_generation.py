"""Document generation: turn plain text/markdown content into a real
downloadable file — .md/.txt need nothing extra, .pdf uses fpdf2 and .docx
uses python-docx (both pure-Python, no system dependency like a headless
LibreOffice or wkhtmltopdf install). Formatting is intentionally simple —
this reads Markdown-ish structure (headings via '#', bullets via '-') well
enough for a generated report or note, not a full CommonMark renderer.
"""

from __future__ import annotations

import io
import re

VALID_FORMATS = {"md", "txt", "pdf", "docx"}
CONTENT_TYPES = {
    "md": "text/markdown",
    "txt": "text/plain",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _plain_bytes(content: str) -> bytes:
    return content.encode("utf-8")


def _to_pdf(content: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    for line in content.splitlines() or [""]:
        heading = re.match(r"^(#{1,3})\s+(.*)", line)
        if heading:
            pdf.set_font("Helvetica", "B", 16 - 2 * (len(heading.group(1)) - 1))
            pdf.multi_cell(0, 8, heading.group(2))
            pdf.set_font("Helvetica", size=11)
        elif not line.strip():
            # A blank line rendered as multi_cell(0, h, " ") leaves fpdf2's
            # cursor in a state where the *next* multi_cell call raises
            # "Not enough horizontal space to render a single character" —
            # a real fpdf2 quirk, not a false alarm. ln() advances the
            # cursor the same visual amount without touching that state.
            pdf.ln(6)
        else:
            text = re.sub(r"^[-*]\s+", "-  ", line)
            # fpdf2's core fonts are Latin-1 only — replace anything outside
            # that range rather than letting the whole generation crash on
            # one smart-quote or emoji the model happened to output.
            safe_text = text.encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 6, safe_text)
    return bytes(pdf.output())


def _to_docx(content: str) -> bytes:
    from docx import Document

    doc = Document()
    for line in content.splitlines() or [""]:
        heading = re.match(r"^(#{1,3})\s+(.*)", line)
        if heading:
            doc.add_heading(heading.group(2), level=len(heading.group(1)))
        elif re.match(r"^[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate(content: str, fmt: str) -> bytes:
    if fmt not in VALID_FORMATS:
        raise ValueError(f"unsupported format: {fmt}")
    if fmt in ("md", "txt"):
        return _plain_bytes(content)
    if fmt == "pdf":
        return _to_pdf(content)
    return _to_docx(content)
